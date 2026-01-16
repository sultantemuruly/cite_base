from fastapi import APIRouter, UploadFile, File, Form
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from docx import Document as DocxDocument
import tempfile
import os

from vectorstore import (
    upload_documents as upload_to_vectorstore,
    get_documents as get_from_vectorstore,
    delete_documents as delete_from_vectorstore,
    update_documents as update_in_vectorstore,
)
from dependencies import SessionDep, CurrentUser
from models import Docs

router = APIRouter(prefix="/documents", tags=["documents"])


def load_document_from_file(file_path: str, file_name: str) -> list[Document]:
    """Convert uploaded file to LangChain Document objects."""
    file_ext = os.path.splitext(file_name)[1].lower()
    documents = []

    try:
        if file_ext not in [".pdf", ".txt", ".docx", ".doc"]:
            raise ValueError(f"Unsupported file type: {file_ext}")
        if file_ext == ".pdf":
            loader = PyPDFLoader(file_path)
            documents = loader.load()
        elif file_ext == ".txt":
            loader = TextLoader(file_path)
            documents = loader.load()
        elif file_ext in [".docx", ".doc"]:
            if file_ext == ".docx":
                doc = DocxDocument(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            else:
                # For .doc files, you might need python-docx or another library
                # For now, we'll try to read as text
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()

            documents = [Document(page_content=text, metadata={"source": file_name})]

        # Add file source to metadata for all documents
        for doc in documents:
            doc.metadata["source"] = file_name

        return documents
    except Exception as e:
        raise Exception(f"Failed to load document {file_name}: {str(e)}")


@router.get("/status")
def get_documents_status():
    return {"status": "Documents route is working"}


@router.get("/get")
def get_documents(ids: list[str]):
    try:
        documents = get_from_vectorstore(ids)
        docs = [doc.dict() for doc in documents]
        return {"status": "success", "documents": docs}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@router.post("/upload")
async def upload_document(
    db_session: SessionDep,
    current_user: CurrentUser,
    title: str = Form(...),
    file: UploadFile = File(...),
):
    try:
        documents = []

        # Create a temporary file
        with tempfile.NamedTemporaryFile(
            delete=False, suffix=os.path.splitext(file.filename)[1]
        ) as tmp_file:
            content = await file.read()
            tmp_file.write(content)
            tmp_file.flush()
            tmp_path = tmp_file.name

        try:
            # Convert file to LangChain Documents
            documents = load_document_from_file(tmp_path, file.filename)
        finally:
            # Clean up temporary file
            os.unlink(tmp_path)

        if not documents:
            return {
                "status": "error",
                "message": "No documents could be extracted from the uploaded files",
            }

        # Upload documents to vector store
        uuids = upload_to_vectorstore(documents)
        
        # Save document metadata to db
        doc = Docs(title=title, user_id=current_user.id, document_uuids=uuids)
        db_session.add(doc)
        db_session.commit()
        db_session.refresh(doc)

        return {"status": "success", "document_id": doc.id, "chunk_ids": uuids, "chunk_count": len(documents)}
    except Exception as e:
        db_session.rollback()
        return {"status": "error", "message": str(e)}


@router.delete("/delete")
def delete_documents(ids: list[str]):
    try:
        delete_from_vectorstore(ids)
        return {"status": "success"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@router.put("/update")
def update_documents(documents: list[Document], ids: list[str]):
    try:
        update_in_vectorstore(documents, ids)
        return {"status": "success"}
    except Exception as e:
        return {"status": "error", "message": str(e)}
